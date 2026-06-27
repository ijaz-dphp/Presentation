import os
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_cv(output_path):
    doc = Document()

    # Name and Contact Info
    name = doc.add_heading('ADNAN RAUF', 0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('Email: adnankhan7847744@gmail.com | Phone: +923042205442\n')
    contact.add_run('LinkedIn: [Add LinkedIn URL] | Location: Bahawalpur, Pakistan')

    # Objective / Professional Summary
    doc.add_heading('Objective & Research Interests', level=1)
    doc.add_paragraph('Highly motivated IT and Administrative professional with over 12 years of experience in the public sector (Combined Military Hospital). Currently pursuing an MS in Public Administration, seeking a fully-funded PhD position starting in Fall 2027. Keenly interested in the intersection of digital governance, public sector management, health administration, and information technology systems in public institutions.')

    # Education
    doc.add_heading('Education', level=1)
    edu1 = doc.add_paragraph()
    edu1.add_run('Master of Science (MS) in Public Administration').bold = True
    edu1.add_run('\n[University Name], [Location]')
    edu1.add_run('\nExpected Graduation: September 2027')
    edu1.add_run('\n• Thesis Topic: [Insert your MS Thesis topic here, e.g., Digital Transformation in Public Healthcare Administration]')

    edu2 = doc.add_paragraph()
    edu2.add_run('\nBachelor\'s Degree in [Insert Subject, e.g., Information Technology / Public Administration]').bold = True
    edu2.add_run('\n[University Name], [Location]')
    edu2.add_run('\nGraduation Year: [Insert Year]')

    # Research Experience & Academic Projects
    doc.add_heading('Research Experience & Academic Projects', level=1)
    doc.add_paragraph('• Currently conducting graduate-level research for MS Thesis in Public Administration.')
    doc.add_paragraph('• Skilled in data analysis and quality management research, specifically applying ISO 9001:2015 frameworks within public healthcare settings.')

    # Professional Experience
    doc.add_heading('Professional & Administrative Experience', level=1)

    exp1 = doc.add_paragraph()
    exp1.add_run('IT Administrator | Combined Military Hospital, Bahawalpur').bold = True
    exp1.add_run('\nJune 2022 – Present')
    exp1.add_run('\n• Manage and maintain hospital IT networks and databases, ensuring seamless operational support for public health administration.')
    exp1.add_run('\n• Oversee the Hospital Management System (HMS), supporting large-scale public service delivery.')

    exp2 = doc.add_paragraph()
    exp2.add_run('\nDatabase System Administrator | Combined Military Hospital, Bahawalpur').bold = True
    exp2.add_run('\nAugust 2015 – June 2022')
    exp2.add_run('\n• Administered the primary database systems for the hospital, ensuring data integrity and security.')
    exp2.add_run('\n• Led the documentation, training, and implementation of ISO SOPs (Quality Management System ISO 9001:2008 and 9001:2015) across hospital staff.')

    exp3 = doc.add_paragraph()
    exp3.add_run('\nAdministrative Staff | Ashraf Sugar Mills Limited, Bahawalpur').bold = True
    exp3.add_run('\n[Insert Dates]')
    exp3.add_run('\n• Provided extensive administrative support, coordinating between departments and streamlining workflow.')

    # Certifications
    doc.add_heading('Certifications & Technical Skills', level=1)
    doc.add_paragraph('• Fundamentals of Digital Marketing – Google.com (40 hours)')
    doc.add_paragraph('• Computer Network Administration (CISCO) – The Islamia University of Bahawalpur')
    doc.add_paragraph('• Freelancing, WordPress, Digital Marketing, SEO, E-Commerce, Creative Writing – Digiskills.pk (Ministry of IT, Pakistan)')
    doc.add_paragraph('• Web and Graphic Designing (Adobe Illustrator/Photoshop) – Vocational Training Institute of Bahawalpur')
    doc.add_paragraph('• Amazon Virtual Assistant – Enablers Pakistan')

    # Additional Information
    doc.add_heading('Additional Information', level=1)
    doc.add_paragraph('• Languages: English (Proficient), Urdu (Native), French (Basic)')
    doc.add_paragraph('• Hobbies & Interests: Reading, traveling, continuous learning, and digital innovation in public spaces.')

    # References
    doc.add_heading('References', level=1)
    ref = doc.add_paragraph()
    ref.add_run('• Brigadier Junaid Masud, Director Medical Services, Headquarters 31 Corps\n')
    ref.add_run('• Brigadier Imran Yousuf, Vice Principal KIMS Karachi\n')
    ref.add_run('• Colonel Nadeem Akhtar, Deputy Commandant, CMH Bahawalpur\n')
    ref.add_run('(Contact details available upon request)')

    doc.save(output_path)
    print(f"CV generated successfully at {output_path}")

if __name__ == '__main__':
    # Use relative path or take from args
    output_file = sys.argv[1] if len(sys.argv) > 1 else os.path.join('PhD_Application_Prep', 'Adnan_Rauf_Academic_CV.docx')
    generate_cv(output_file)
